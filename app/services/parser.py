import os
import pypdf
import docx
import logging

logger = logging.getLogger(__name__)

class ParsingError(Exception):
    """Custom exception raised when document parsing fails."""
    pass

def parse_pdf(local_path: str) -> list[dict]:
    """
    Parses a PDF file page by page.
    Returns a list of dictionaries with text content and page metadata.
    """
    try:
        reader = pypdf.PdfReader(local_path)
        blocks = []
        
        # Guard against encrypted files
        if reader.is_encrypted:
            try:
                # Try decrypting with empty password
                reader.decrypt("")
            except Exception:
                raise ParsingError("PDF is encrypted and cannot be decrypted.")

        for page_idx, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
                text = text.strip()
                if text:
                    blocks.append({
                        "content": text,
                        "page": page_idx + 1,  # 1-indexed page
                        "position": f"page_{page_idx + 1}"
                    })
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_idx + 1}: {str(e)}")
                # Continue extracting other pages if possible
        
        if not blocks:
            raise ParsingError("No text content could be extracted from this PDF.")
            
        return blocks
    except ParsingError:
        raise
    except Exception as e:
        raise ParsingError(f"Failed to parse PDF file: {str(e)}")

def parse_docx(local_path: str) -> list[dict]:
    """
    Parses a Word file paragraph by paragraph.
    Returns a list of dictionaries with text content and paragraph position.
    """
    try:
        doc = docx.Document(local_path)
        blocks = []
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                blocks.append({
                    "content": text,
                    "page": 1,  # Word documents do not have definitive page concepts without rendering
                    "position": f"paragraph_{para_idx + 1}"
                })
                
        # Word docs can also have tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    blocks.append({
                        "content": row_text,
                        "page": 1,
                        "position": f"table_{table_idx + 1}_row_{row_idx + 1}"
                    })

        if not blocks:
            raise ParsingError("No text content could be extracted from this DOCX file.")
            
        return blocks
    except Exception as e:
        raise ParsingError(f"Failed to parse DOCX file: {str(e)}")

def parse_txt(local_path: str) -> list[dict]:
    """
    Parses a plain text file.
    Attempts UTF-8, then falls back to Latin-1 and CP1252.
    """
    encodings = ["utf-8", "latin-1", "cp1252"]
    text = None

    if not os.path.exists(local_path):
        raise ParsingError(f"File not found locally: {local_path}")

    for encoding in encodings:
        try:
            with open(local_path, "r", encoding=encoding) as f:
                text = f.read()
            break  # Successfully read file
        except UnicodeDecodeError:
            continue
            
    if text is None:
        raise ParsingError("Failed to decode text file. Unsupported encoding.")
        
    text = text.strip()
    if not text:
        raise ParsingError("Plain text file is empty.")

    # Split plain text into paragraph blocks for easier processing
    paragraphs = text.split("\n\n")
    blocks = []
    
    for idx, para in enumerate(paragraphs):
        para_text = para.strip()
        if para_text:
            blocks.append({
                "content": para_text,
                "page": 1,
                "position": f"paragraph_{idx + 1}"
            })
            
    return blocks

def parse_document(local_path: str, content_type: str) -> list[dict]:
    """
    Dispatcher to parse document based on content type.
    """
    if content_type == "application/pdf":
        return parse_pdf(local_path)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return parse_docx(local_path)
    elif content_type == "text/plain":
        return parse_txt(local_path)
    else:
        raise ParsingError(f"Unsupported content type for parsing: {content_type}")
